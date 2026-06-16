#!/usr/bin/env python3
"""
build_docx.py -- generate an editable Word (.docx) version of the paper.

Converts main.tex -> main.docx via pandoc. Because the manuscript uses a
Science-style two-column title block, custom macros (\arch, \archentry,
\bbar, \rtp), tcolorbox story hands, and LaTeX \ref cross-references --
none of which survive a naive pandoc run -- this script first rewrites the
source into a pandoc-friendly form:

  * the \twocolumn[...] title block becomes \title/\author + abstract;
  * \input table files are inlined;
  * \arch{Name} is expanded to the archetype's brand hex color (bold);
  * \archentry / \bbar / tcolorbox are flattened to plain text + headings;
  * every figure/table/equation is numbered and its \ref{} is resolved to
    that number, with "Figure N." / "Table N." prefixed onto each caption;
  * citations are rendered numbered (IEEE) with a reference list.

Requires: pandoc (with citeproc). Run from anywhere:
    python3 build_docx.py
Output: main.docx (next to main.tex).
"""
import re, pathlib, subprocess, sys, tempfile

MAN = pathlib.Path(__file__).resolve().parent
src = (MAN / 'main.tex').read_text()

# ---------------------------------------------------------------- 1. inline \input
src = re.sub(r'\\input\{([^}]+)\}',
             lambda m: (MAN / m.group(1)).resolve().read_text(), src)

# ---------------------------------------------------------------- 2. title block
TITLE = ("Trust Dynamics in Multi-Agent Strategic Interaction: A Simulation "
         "Study of Bayesian Reputation Systems in 8-Player Limit Texas Hold'em")
AUTHOR = (r"Rachit Agrawal\\ \normalsize Independent research, 2025--2026.\\ "
          r"\normalsize Correspondence: rachit.agrawal@sahyadrischool.org")
ABSTRACT = (
 r"WHAT IF BEING TRUSTWORTHY COST YOU SOMETHING? Reputation systems (online "
 r"marketplaces, credit scoring and social platforms) reward ``good'' "
 r"behavior---however, there is a well-known paradox among economists. A totally "
 r"transparent and cooperative player will get less than a somewhat unreliable "
 r"player. We placed that paradox into a petri-dish. Eight poker-players played "
 r"tens of thousands of hands. Each player silently profiled all other players' "
 r"styles, and updated those beliefs at each hand. Then we swapped the minds of "
 r"the agents across four different levels of complexity---from strict rule "
 r"followers to large language models capable of reasoning about what their "
 r"opponents would do, while keeping both the game and the trust model constant. "
 r"And so, it happened again. The best reputations were the least rewarded, and "
 r"no amount of optimization changed that. It wasn't until the agents were able "
 r"to reason about the reputation system that some relief began to occur. The "
 r"result is a compact, completely replicable laboratory for exploring the very "
 r"important question---when does a good reputation become a bad thing?")

_abs = '\\begin{abstract}\n' + ABSTRACT + '\n\\end{abstract}'
src = re.sub(r'\\twocolumn\[.*?\\thispagestyle\{firstpage\}',
             lambda m: _abs, src, count=1, flags=re.S)

preamble_inject = ('\\title{' + TITLE + '}\n\\author{' + AUTHOR + '}\n\\date{}\n'
                   '\\begin{document}')
src = src.replace('\\begin{document}', preamble_inject, 1)

# ---------------------------------------------------------------- 3. tcolorbox -> heading + text
src = re.sub(r'\\begin\{tcolorbox\}\[storyhand=\{(.*?)\}\]',
             r'\\par\\medskip\\noindent\\textbf{\1}\\par\\medskip', src, flags=re.S)
src = src.replace(r'\end{tcolorbox}', r'\par\medskip')

# ---------------------------------------------------------------- 4. number map (document order)
labmap, fig, tab, eq, cur = {}, 0, 0, 0, None
for m in re.finditer(r'\\begin\{(figure\*?|table\*?|equation)\}|\\label\{([^}]+)\}', src):
    env, lab = m.group(1), m.group(2)
    if env:
        if env.startswith('figure'): fig += 1; cur = ('fig', fig)
        elif env.startswith('table'): tab += 1; cur = ('tab', tab)
        else: eq += 1; cur = ('eq', eq)
    elif lab and cur:
        labmap[lab] = cur; cur = None
labmap['sec:results'] = ('sec', '4')      # pandoc --number-sections numbering
labmap['sec:archetypes'] = ('sec', '3.2')
labmap['sec:posterior'] = ('sec', '3.3')

# ---------------------------------------------------------------- 5. prefix captions
def _float_repl(m):
    whole, body = m.group(0), m.group(2)
    lm = re.search(r'\\label\{([^}]+)\}', body)
    if lm and lm.group(1) in labmap:
        typ, num = labmap[lm.group(1)]
        kind = 'Figure' if typ == 'fig' else 'Table'
        return re.sub(r'\\caption\{',
                      r'\\caption{\\textbf{' + f'{kind}~{num}.' + r'}~', whole, count=1)
    return whole
src = re.sub(r'\\begin\{(figure\*?|table\*?)\}(.*?)\\end\{\1\}', _float_repl, src, flags=re.S)

# ---------------------------------------------------------------- 6. equations: append (N)
def _eq_repl(m):
    body = m.group(1)
    lm = re.search(r'\\label\{([^}]+)\}', body)
    num = labmap.get(lm.group(1), ('eq', '?'))[1] if lm else '?'
    body = re.sub(r'\s*\\label\{[^}]+\}', '', body).rstrip()
    return r'\begin{equation}' + body + r'\qquad(' + str(num) + ')\n' + r'\end{equation}'
src = re.sub(r'\\begin\{equation\}(.*?)\\end\{equation\}', _eq_repl, src, flags=re.S)

# ---------------------------------------------------------------- 7. resolve \ref, drop \label
src = re.sub(r'\\ref\{([^}]+)\}', lambda m: str(labmap.get(m.group(1), ('?', '?'))[1]), src)
src = re.sub(r'\\label\{[^}]+\}', '', src)

# ---------------------------------------------------------------- 8. archetype brand colors
HEX = {'Oracle':'A8B5C3','Sentinel':'5C8159','Firestorm':'E75D3C','Wall':'A2A7A9',
       'Phantom':'836F91','Predator':'B93C41','Mirror':'CDD3D8','Judge':'3F4F7C'}
src = re.sub(r'\\arch\{(\w+)\}',
             lambda m: r'\textbf{' + m.group(1) + r'}', src)
# unwrap drafted-prose highlight \W{...} -> plain text for the Word build
# (handles one level of nested braces, e.g. \W{... \citep{key} ...})
src = re.sub(r'\\W\{((?:[^{}]|\{[^{}]*\})*)\}', r'\1', src)

# ---------------------------------------------------------------- 9. flatten custom macros
renews = (r'\renewcommand{\bbar}{ \textbar{} }' + '\n'
          r'\renewcommand{\archentry}[5]{\par\medskip\noindent\textbf{#1} \textbar{} '
          r'\textit{#2} \textbar{} honesty #3. #4\ (#5)\par}' + '\n')
src = src.replace(preamble_inject, renews + preamble_inject, 1)

# ---------------------------------------------------------------- 10. image widths fit a page
src = re.sub(r'width=\\linewidth', 'width=14cm', src)
src = re.sub(r'width=0\.9\d*\\textwidth', 'width=15cm', src)

# ---------------------------------------------------------------- 11. post-process
def _postprocess(path):
    """Recolor archetype names (pandoc drops \\textcolor) and add the running
    header / footer as real Word header & footer objects. Best-effort: if
    python-docx is unavailable the docx is still valid, just without colors
    and running heads."""
    try:
        from docx import Document
        from docx.shared import RGBColor, Inches, Pt
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print('  (python-docx not installed; skipping colors + headers/footers)')
        return
    doc = Document(str(path))

    def _paras(container):
        for p in container.paragraphs:
            yield p
        for t in container.tables:
            for row in t.rows:
                for cell in row.cells:
                    yield from _paras(cell)

    # archetype names are rendered in plain bold (no per-archetype colour, v6+)

    def _pagefield(par):
        rr = par.add_run()
        for kind, val in (('begin', None), ('instr', ' PAGE '), ('end', None)):
            if kind == 'instr':
                e = OxmlElement('w:instrText'); e.set(qn('xml:space'), 'preserve'); e.text = val
            else:
                e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), kind)
            rr._r.append(e)

    for sec in doc.sections:
        sec.different_first_page_header_footer = False
        sec.header.is_linked_to_previous = False
        hp = sec.header.paragraphs[0]; hp.text = ''; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rh = hp.add_run('RESEARCH  |  RESEARCH ARTICLE')
        rh.bold = True; rh.font.size = Pt(8); rh.font.color.rgb = RGBColor.from_string('C8102E')
        sec.footer.is_linked_to_previous = False
        fp = sec.footer.paragraphs[0]; fp.text = ''
        fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        r1 = fp.add_run('Agrawal (2026)'); r1.font.size = Pt(8)
        fp.add_run('\t'); _pagefield(fp)

    doc.save(str(path))
    print('  post-processed: running header/footer added (archetype names plain bold)')

# ---------------------------------------------------------------- write + pandoc
tmp = pathlib.Path(tempfile.mkdtemp()) / 'main_pandoc.tex'
tmp.write_text(src)
out = MAN / 'main.docx'
cmd = ['pandoc', str(tmp), '-o', str(out), '--citeproc',
       '--csl=' + str(MAN / 'ieee.csl'), '--bibliography=' + str(MAN / 'references.bib'),
       '--number-sections', '--resource-path=' + str(MAN) + ':' + str(MAN.parent / 'figures')]
print('pandoc:', ' '.join(cmd))
r = subprocess.run(cmd)
if r.returncode == 0:
    _postprocess(out)
    print(f'wrote {out}  ({fig} figures, {tab} tables, {eq} equations)')
sys.exit(r.returncode)
